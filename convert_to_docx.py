#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Document Converter (Compact Version)
Markdownファイルをコンパクトで読みやすいWord文書(.docx)に変換します
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """段落にハイパーリンクを追加"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # ハイパーリンクスタイル
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def process_inline_formatting(paragraph, text):
    """インライン書式（太字、イタリック、リンク等）を処理"""
    # リンク処理 [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
    matches = list(re.finditer(link_pattern, text))

    if matches:
        last_end = 0
        for match in matches:
            # リンク前のテキスト
            if match.start() > last_end:
                before_text = text[last_end:match.start()]
                process_text_formatting(paragraph, before_text)

            # リンク
            link_text = match.group(1)
            link_url = match.group(2)
            add_hyperlink(paragraph, link_text, link_url)

            last_end = match.end()

        # リンク後のテキスト
        if last_end < len(text):
            after_text = text[last_end:]
            process_text_formatting(paragraph, after_text)
    else:
        process_text_formatting(paragraph, text)

def process_text_formatting(paragraph, text):
    """太字、イタリック等のテキスト書式を処理"""
    # **bold** または __bold__
    bold_pattern = r'\*\*([^\*]+)\*\*|__([^_]+)__'
    # `code`
    code_pattern = r'`([^`]+)`'

    parts = re.split(r'(\*\*[^\*]+\*\*|__[^_]+__|`[^`]+`)', text)

    for part in parts:
        if not part:
            continue

        # 太字
        if re.match(r'^\*\*.*\*\*$|^__.*__$', part):
            clean_text = re.sub(r'^\*\*|\*\*$|^__|__$', '', part)
            run = paragraph.add_run(clean_text)
            run.bold = True
        # コード
        elif re.match(r'^`.*`$', part):
            clean_text = re.sub(r'^`|`$', '', part)
            run = paragraph.add_run(clean_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)

def convert_markdown_to_docx(md_file, docx_file):
    """MarkdownファイルをWord文書に変換（コンパクト版）"""
    # Markdownファイルを読み込み
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Word文書を作成
    doc = Document()

    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Meiryo'
    font.size = Pt(10)  # 10ptに縮小

    # 段落間隔を縮小
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(3)  # 3ptに縮小
    paragraph_format.line_spacing = 1.15  # 行間を縮小

    in_code_block = False
    code_lines = []
    in_list = False
    skip_next_empty = False
    consecutive_empty = 0

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # コードブロック開始/終了
        if line.startswith('```'):
            if in_code_block:
                # コードブロック終了 - 内容を追加（コンパクト）
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8)  # さらに小さく
                # 背景色（グレー）
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')
                p._p.get_or_add_pPr().append(shading_elm)
                # 間隔を縮小
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.25)

                code_lines = []
                in_code_block = False
                skip_next_empty = True
            else:
                in_code_block = True
            i += 1
            continue

        # コードブロック内
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 空行（連続する空行を1つにまとめる）
        if not line.strip():
            consecutive_empty += 1
            if consecutive_empty == 1 and not skip_next_empty:
                # 空行は段落間隔で表現（実際には追加しない）
                pass
            skip_next_empty = False
            if in_list:
                in_list = False
            i += 1
            continue

        consecutive_empty = 0

        # 見出し
        if line.startswith('#'):
            in_list = False
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, level=1)
                p.runs[0].font.size = Pt(16)  # 縮小
                p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            elif level == 2:
                p = doc.add_heading(text, level=2)
                p.runs[0].font.size = Pt(13)  # 縮小
                p.runs[0].font.color.rgb = RGBColor(0, 102, 153)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
            elif level == 3:
                p = doc.add_heading(text, level=3)
                p.runs[0].font.size = Pt(11)  # 縮小
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
            else:
                p = doc.add_paragraph(text)
                p.runs[0].bold = True
                p.paragraph_format.space_after = Pt(3)

        # 水平線（省略 - 見出しで十分）
        elif line.startswith('---') or line.startswith('***'):
            in_list = False
            # 水平線は省略してスペースのみ
            pass

        # テーブル
        elif '|' in line and not line.strip().startswith('|'):
            in_list = False
            # テーブル行を収集
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip())
                i += 1
            i -= 1

            # テーブルをパース
            if len(table_lines) >= 2:
                # ヘッダー行
                header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # データ行（区切り行をスキップ）
                data_rows = []
                for table_line in table_lines[2:]:
                    if table_line.strip().startswith('|') and not all(c in '-:|' for c in table_line.replace(' ', '')):
                        row = [cell.strip() for cell in table_line.split('|')[1:-1]]
                        data_rows.append(row)

                # Word文書にテーブル追加
                if data_rows:
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
                    table.style = 'Light Grid Accent 1'

                    # テーブル全体のフォントサイズを縮小
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)

                    # ヘッダー
                    for j, cell_text in enumerate(header):
                        cell = table.rows[0].cells[j]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(9)

                    # データ行
                    for row_idx, row_data in enumerate(data_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx + 1].cells):
                                cell = table.rows[row_idx + 1].cells[col_idx]
                                cell.text = cell_text
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)

                    # テーブル後の余白を縮小
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.paragraph_format.space_after = Pt(6)

        # 順序付きリスト
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text)
            p.paragraph_format.space_after = Pt(2)  # リスト項目間を縮小
            for run in p.runs:
                run.font.size = Pt(10)
            in_list = True

        # 順序なしリスト
        elif re.match(r'^[-*+]\s', line):
            # インデントレベルを検出
            indent = len(line) - len(line.lstrip())
            text = re.sub(r'^[-*+]\s', '', line.lstrip())

            p = doc.add_paragraph(style='List Bullet')
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.2 * (indent // 2 + 1))  # 縮小
            process_inline_formatting(p, text)
            p.paragraph_format.space_after = Pt(2)  # リスト項目間を縮小
            for run in p.runs:
                run.font.size = Pt(10)
            in_list = True

        # 通常の段落
        else:
            if in_list:
                in_list = False
            p = doc.add_paragraph()
            process_inline_formatting(p, line)
            p.paragraph_format.space_after = Pt(3)  # 段落間を縮小
            for run in p.runs:
                run.font.size = Pt(10)

        i += 1

    # Word文書を保存
    doc.save(docx_file)
    print(f'✓ Created: {docx_file}')

def main():
    """メイン処理"""
    import os

    consulting_dir = 'consulting'

    # 変換対象ファイル
    files_to_convert = [
        ('consulting_agenda.md', 'consulting_agenda.docx'),
        ('genspark_prompt_corrected.md', 'genspark_prompt_corrected.docx'),
        ('agenda_import_guide.md', 'agenda_import_guide.docx'),
        ('kickstarter_demo_guide.md', 'kickstarter_demo_guide.docx'),
    ]

    print('Converting Markdown files to Word documents (Compact version)...\n')

    for md_file, docx_file in files_to_convert:
        md_path = os.path.join(consulting_dir, md_file)
        docx_path = os.path.join(consulting_dir, docx_file)

        if os.path.exists(md_path):
            convert_markdown_to_docx(md_path, docx_path)
        else:
            print(f'⚠ Not found: {md_path}')

    print('\n✅ Conversion completed!')
    print(f'\nCompact Word documents are saved in: {consulting_dir}/')
    print('\nChanges from previous version:')
    print('  - Reduced font size (10pt)')
    print('  - Reduced line spacing (1.15)')
    print('  - Reduced paragraph spacing (3pt)')
    print('  - Removed horizontal lines')
    print('  - Smaller tables (9pt)')
    print('  - More compact overall layout')

if __name__ == '__main__':
    main()
